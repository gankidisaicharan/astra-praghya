# ═══════════════════════════════════════════════════════════════════
# Astra v2.1 — Praghya Prakhar
# Supply Chain & Operations + Graduate Management Resume Tailoring Engine
#
# v2.1 changes (May 2026):
#  - Two-track strategy: handles both direct-hire SC roles AND graduate
#    management programmes (Tesco, Kerry, Bank of Ireland, Big 4, etc.)
#  - JD-archetype detector — 7 archetypes drive different framing
#  - Stamp 1G work authorisation surfaced as a positive differentiator
#  - NFQ Level 9 framing for the MSc (Irish-recruiter friendly)
#  - Expanded vocabulary: GRN processing, cycle counting, stock
#    reconciliation, PO verification, distribution centre, 3PL
#    coordination, OTIF awareness, continuous improvement
#  - Industry-bridge logic for pharma / FMCG / logistics / financial
#  - jd_fit_warning surfaced prominently in UI for weak matches
#  - War-story #6 added to cover letter prompt for graduate programmes
#
# Design goals (unchanged):
#  - Simple. No conditional gates that block CV generation.
#  - Truthful. Tailor to JD using ONLY skills/experience she actually has.
#  - Structure-preserving. Every base-resume section is always present.
#  - All her real base skills are kept; JD-relevant skills are added on top.
#  - Education / Certifications / Additional Info: untouched, always rendered.
#  - ATS-friendly output every run.
# ═══════════════════════════════════════════════════════════════════

import streamlit as st
import json
import re
import io
import datetime
from typing import List, Optional
from pydantic import BaseModel, Field
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from xml.sax.saxutils import escape


# ═══════════════════════════════════════════════════════════════════
# 1. CONFIG
# ═══════════════════════════════════════════════════════════════════

PAGE_TITLE = "Astra — Praghya Prakhar"

# Generation model.
# As of May 2026, gemini-3-flash-preview is the recommended free-tier model
# in the Gemini 3 family. It handles structured-JSON output reliably.
#
# Migration notes — if you see errors, swap MODEL string to one of:
#   - "gemini-3.1-flash-lite-preview"  (cheaper, slightly weaker reasoning)
#   - "gemini-2.5-flash"               (stable until June 17, 2026)
#   - "gemini-2.5-pro"                 (best quality, lower rate limits, deprecates June 17, 2026)
GENERATION_MODEL = "gemini-3-flash-preview"

try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
except Exception:
    GOOGLE_API_KEY = ""


# ═══════════════════════════════════════════════════════════════════
# 2. FIXED CANDIDATE FACTS
# These NEVER change between runs, no matter what the JD says.
# ═══════════════════════════════════════════════════════════════════

CANDIDATE_NAME = "Praghya Prakhar"
CANDIDATE_TAGLINE = "Supply Chain & Operations Professional"
CANDIDATE_CONTACT = "Dublin, Ireland | +353 89 263 0034 | praghyaprakhar2012@gmail.com | linkedin.com/in/praghya-prakhar-a9b016209"

# Base skill set — real and verified. These are ALWAYS in the output.
# Each category is a list of skills. The tailoring step may APPEND
# JD-relevant skills, but it cannot remove any of these.
BASE_SKILLS = {
    "Supply Chain Management": [
        "Inventory Control", "Order Fulfilment", "Warehouse Operations",
        "Distribution Centre Operations", "Inbound/Outbound Logistics",
        "Stock Auditing", "Cycle Counting", "Stock Reconciliation",
        "Dispatch Coordination", "Goods Receipt Note (GRN) Processing",
        "Purchase Order (PO) Verification", "3PL Coordination",
    ],
    "Operations & Process Improvement": [
        "Process Standardisation", "Operational Efficiency",
        "Quality Assurance", "SOP Development", "KPI Monitoring",
        "Service Level Monitoring", "Continuous Improvement",
    ],
    "ERP & Software": [
        "SAP (Inventory & SCM Modules)", "Oracle Fusion Cloud SCM",
        "Microsoft Excel", "Microsoft Word", "Google Colab",
    ],
    "Leadership & Coordination": [
        "Team Training & Mentoring", "Cross-Functional Coordination",
        "Vendor Liaison", "Stakeholder Communication",
        "POSH Compliance Training",
    ],
    "Analytical & Research": [
        "Data Collection & Analysis", "Survey Design (Google Forms)",
        "Report Preparation", "Published Research",
    ],
}

# Three real roles. Tailoring rewrites the BULLETS within these,
# but never adds/removes roles or changes company/dates/title.
BASE_EXPERIENCE = [
    {
        "role_title": "Senior Graduate, Operations & Supply Chain",
        "company": "Reliance Retail (Quick Supply Chain Division)",
        "location": "Delhi, India",
        "dates": "Aug 2022 – Dec 2024",
        "responsibilities": [
            "Managed daily distribution centre operations including inbound receipts, outbound dispatch, and stock reconciliation for a fulfilment hub servicing 250+ retail stores and e-commerce orders.",
            "Oversaw inventory accuracy across ~5,000 SKUs through systematic cycle counting and stock audits, maintaining accuracy levels above 97%.",
            "Coordinated cross-functionally with procurement, 3PL logistics partners, and store operations teams to reduce order dispatch delays by ~15%, lifting on-time-in-full performance against service-level targets.",
            "Streamlined inbound shipment processing and Goods Receipt Note (GRN) workflows, cutting average goods-in turnaround time by ~20% through improved staging and documentation procedures.",
            "Trained and mentored a team of 20+ warehouse staff on operational processes, safety protocols, hygiene standards, and POSH compliance, with a focus on onboarding female employees.",
        ],
        "achievements": [
            "Promoted from Graduate Trainee to Senior Graduate within 6 months based on strong operational performance and leadership.",
            "Recognised internally for improving dispatch reliability and warehouse floor discipline across the fulfilment centre.",
        ],
    },
    {
        "role_title": "Logistics Intern",
        "company": "Om Logistics",
        "location": "Delhi, India",
        "dates": "Jun 2021 – Aug 2021",
        "responsibilities": [
            "Tracked and monitored consignment movements across multiple routes, identifying and resolving shipment delays to maintain delivery timelines.",
            "Coordinated with vendors and internal warehouse teams to streamline incoming shipment processing and improve Goods Receipt Note (GRN) accuracy.",
            "Assisted in shipment clearance procedures, reducing documentation bottlenecks and improving clearance turnaround by ~10%.",
            "Verified incoming inventory against purchase orders (PO verification) and maintained accurate records of ~200+ weekly consignments.",
        ],
        "achievements": [],
    },
    {
        "role_title": "Operations Intern",
        "company": "Shubh Consultants & Technocrats LLP",
        "location": "Delhi, India",
        "dates": "Jun 2022 – Jul 2022",
        "responsibilities": [
            "Supported day-to-day project coordination and documentation activities across multiple consulting engagements.",
            "Maintained project records, performed data entry, and organised operational information to ensure smooth workflow across teams.",
            "Gained practical exposure to structured project management processes and cross-functional team coordination in a professional consulting environment.",
        ],
        "achievements": [],
    },
]

# Education, certifications, additional info — NEVER tailored.
# Always rendered as-is.
BASE_EDUCATION = [
    {
        "degree": "MSc in Management (Strategy), NFQ Level 9",
        "institution": "Dublin City University (DCU), Dublin, Ireland",
        "dates": "Jan 2025 – Mar 2026",
        "grade": "Grade: 2:1",
        "extra": "Dissertation: Impact of Social Media on Decision-Making and Emotional Well-Being. Primary data collection via surveys, analysis using Excel and Google Colab.",
    },
    {
        "degree": "BBA in Logistics & Supply Chain Management",
        "institution": "Galgotias University, Greater Noida, India",
        "dates": "May 2019 – May 2022",
        "grade": "GPA: 9.2/10 | Silver Medallist (Top 2)",
        "extra": "Published research paper on supply chain management challenges, analysing operational inefficiencies and proposing improvement frameworks.",
    },
]

BASE_CERTIFICATIONS = [
    "Oracle Fusion Cloud Applications SCM Process Essentials Certified (Rel 1), Oracle University, November 2025",
    "Forage Virtual Experience Programme: Client Analysis, Sustainability Solutions & Fitment Matrix Presentation, September 2025",
    "Processes in SAP S/4HANA Extended Warehouse Management (EWM) (In Progress)",
]

BASE_ADDITIONAL_INFO = [
    "Languages: English (Fluent), Hindi (Native)",
    "Awards: DCU Scholarship Recipient (€2,000); Silver Medal, BBA Graduation (Top 2)",
    "Volunteer & Extracurricular: CRY (NGO) Volunteer; Marketing Club Member; NCC B Certificate Holder",
    "Currently Learning: Six Sigma Green Belt, PMP (Project Management Professional)",
    "Work Authorisation: Eligible for full-time employment in Ireland under Stamp 1G (Third Level Graduate Programme). No employment permit required.",
]


def build_base_resume_text() -> str:
    """Compose Praghya's base resume as plain text from the constants above.
    This is what pre-fills the 'Base Resume' box in the UI so she can see
    and edit what's being sent to the model."""
    lines = []
    lines.append(CANDIDATE_NAME)
    lines.append(CANDIDATE_TAGLINE)
    lines.append(CANDIDATE_CONTACT)
    lines.append("")

    lines.append("Professional Profile")
    lines.append(
        "Operations and supply chain professional with over 2 years of experience "
        "in warehouse management, inventory control, and logistics coordination "
        "within large-scale retail environments. Held a Senior Graduate role at "
        "Reliance Retail, one of India's largest retail conglomerates, overseeing "
        "end-to-end fulfilment operations across 250+ stores and maintaining 97% "
        "inventory accuracy across ~5,000 SKUs. Holds an MSc in Management "
        "(Strategy), NFQ Level 9, from Dublin City University and a BBA in "
        "Logistics & Supply Chain Management (9.2/10 GPA, Silver Medallist). "
        "Certified in Oracle Fusion Cloud SCM. Eligible for full-time employment "
        "in Ireland under Stamp 1G, with no employment permit required. Open to "
        "supply chain coordinator, operations, logistics, and graduate management "
        "roles in the Irish market."
    )
    lines.append("")

    lines.append("Key Skills / Tools & Technologies")
    for cat, skill_list in BASE_SKILLS.items():
        lines.append(f"- {cat}: {', '.join(skill_list)}")
    lines.append("")

    lines.append("Professional Experience")
    for role in BASE_EXPERIENCE:
        header = f"{role['role_title']} | {role['company']} | {role['location']} | {role['dates']}"
        lines.append(header)
        for r in role["responsibilities"]:
            lines.append(f"- {r}")
        if role["achievements"]:
            lines.append("Achievements:")
            for a in role["achievements"]:
                lines.append(f"- {a}")
        lines.append("")

    lines.append("Education")
    for edu in BASE_EDUCATION:
        lines.append(edu["degree"])
        lines.append(f"{edu['institution']} | {edu['dates']} | {edu['grade']}")
        if edu.get("extra"):
            lines.append(edu["extra"])
        lines.append("")

    lines.append("Certifications")
    for c in BASE_CERTIFICATIONS:
        lines.append(f"- {c}")
    lines.append("")

    lines.append("Additional Information")
    for a in BASE_ADDITIONAL_INFO:
        lines.append(f"- {a}")

    return "\n".join(lines)


PRAGHYA_BASE_RESUME = build_base_resume_text()


# ═══════════════════════════════════════════════════════════════════
# 3. SAFETY: BANNED SKILLS (Praghya does NOT have these)
# Used as a final scrub on tailored output. We don't fail — we just strip.
# ═══════════════════════════════════════════════════════════════════

BANNED_SKILLS = {
    # Programming
    "python", "sql", "r programming", "javascript", "java", "c++", "c#", "golang", "rust",
    # BI / Viz
    "power bi", "powerbi", "tableau", "looker", "qlik", "qlikview", "data studio",
    # Advanced Excel
    "vba", "macros", "advanced excel",
    # Cloud / DBs
    "aws", "azure", "gcp", "google cloud platform",
    "postgresql", "mysql", "mongodb", "redis", "snowflake", "bigquery",
    # Data Science / ML
    "machine learning", "deep learning", "data science", "nlp",
    "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy",
    # Certs she's only "in progress" with — must not appear as completed
    "six sigma certified", "six sigma black belt", "six sigma green belt certified",
    "pmp certified", "prince2 certified",
    # DevOps
    "docker", "kubernetes", "ci/cd", "terraform",
    # Frontend frameworks
    "react", "angular", "node.js", "vue.js",
}


# ═══════════════════════════════════════════════════════════════════
# 4. PROMPT — single, focused, no over-engineering
# ═══════════════════════════════════════════════════════════════════

ASTRA_PROMPT = """You are Astra, a resume tailoring engine for Praghya Prakhar, a Supply Chain & Operations professional based in Dublin, Ireland.

Your job: take Praghya's BASE RESUME and the JOB DESCRIPTION, and produce a tailored version that mirrors the JD's language and priorities, while staying 100% truthful.

═══ HARD RULE: NO EM DASHES ANYWHERE IN YOUR OUTPUT ═══

This is non-negotiable. Em dashes (—, U+2014) are a leading recruiter signal of AI-generated writing in 2026. They MUST NOT appear in:
- The summary
- Any experience bullet
- The candidate_title
- The jd_fit_warning
- Any other field

Replace every em dash with one of: a comma, a full stop with sentence break, a colon, parentheses, or simply rephrase the sentence to not need one. Hyphens inside compound words (e.g. "cross-functional", "on-time-in-full") and en dashes inside date ranges (e.g. "Aug 2022 – Dec 2024") are fine and not affected by this rule.

Examples of em-dash phrasing to AVOID and how to fix:
- BAD:  "She is on Stamp 1G — no permit required."
  GOOD: "She is on Stamp 1G, with no permit required."
- BAD:  "MSc in Management (Strategy) — NFQ Level 9, DCU"
  GOOD: "MSc in Management (Strategy), NFQ Level 9, DCU"
- BAD:  "2 years at Reliance Retail — managing fulfilment for 250+ stores"
  GOOD: "2 years at Reliance Retail managing fulfilment for 250+ stores"

═══ HARD RULE: NO PRONOUNS IN THE SUMMARY ═══

The professional profile / summary section MUST be written in pronoun-free, implied-subject style. This is the dominant Irish/UK CV convention.

DO NOT use any of these:
- First-person: "I", "I am", "I bring", "my experience"
- Third-person: "She", "She brings", "She is", "her experience", "Praghya is", "Praghya brings"
- Possessives: "my MSc", "her promotion"

INSTEAD: Drop the subject entirely. The reader knows it's about the candidate.

Examples of bad and good summary openings:
- BAD:  "I am an operations professional with over 2 years..."
- BAD:  "She brings real operational depth from 2 years..."
- BAD:  "Praghya was promoted from Graduate Trainee..."
- GOOD: "Operations professional with over 2 years..."
- GOOD: "Brings real operational depth from 2 years..."  (verb-led, no pronoun)
- GOOD: "Promoted from Graduate Trainee to Senior Graduate within 6 months..."

This applies to the summary ONLY. Cover letters use first-person ("I", "my") which is normal letter convention. Bullets use action verbs (Managed, Coordinated, Reduced) which are already pronoun-free.

═══ CANDIDATE FACTS YOU CANNOT INVENT ═══

Praghya's REAL skills (safe to claim):
- ERP / Software: SAP (Inventory & SCM Modules, daily use at Reliance), Oracle Fusion Cloud SCM (certified Nov 2025), Microsoft Excel (basic-intermediate), Microsoft Word, Google Colab (basic).
- Supply Chain: Inventory control, order fulfilment, warehouse / distribution centre operations, inbound/outbound logistics flows, stock auditing, cycle counting, stock reconciliation, dispatch coordination, Goods Receipt Note (GRN) processing, Purchase Order (PO) verification, 3PL coordination, service level monitoring, OTIF awareness.
- Operations: Process standardisation, SOP development, KPI monitoring, quality assurance, continuous improvement.
- Leadership: Team training & mentoring, cross-functional coordination, vendor liaison, stakeholder communication, POSH compliance training.
- Analytical: Data collection via surveys (Google Forms), report preparation, published research on SCM.
- Education: MSc in Management (Strategy), NFQ Level 9, from Dublin City University, Grade 2:1. BBA in Logistics & Supply Chain Management with 9.2/10 GPA and Silver Medal (Top 2 of class).
- Work Authorisation: Stamp 1G (Third Level Graduate Programme). Eligible for full-time employment in Ireland, no employment permit required.

Praghya does NOT have, and you must NEVER claim she does:
- Programming: Python, SQL, R, JavaScript, Java, C++.
- BI tools: Power BI, Tableau, Looker, Qlik.
- Advanced Excel: VBA, macros, complex pivot/array formulas.
- Cloud: AWS, Azure, GCP. Databases: PostgreSQL, MySQL, MongoDB, Snowflake.
- Data science / ML / deep learning.
- Six Sigma certification or PMP certification (currently learning, NOT certified).
- S&OP experience, GDP/GMP regulatory experience (only awareness via coursework, do not claim hands-on).
- Any Irish professional work experience.

If the JD demands Python / SQL / Power BI / S&OP / GDP as must-have, do NOT add them. Honesty is non-negotiable.

═══ STEP 1: DETECT THE JD ARCHETYPE ═══

Before tailoring, classify the JD into ONE of these seven archetypes. This drives how you frame the summary, bullets, and skills.

A. supply_chain_coordinator
   Signals: titles like "Supply Chain Coordinator / Specialist / Assistant / Logistics Coordinator / Inventory Coordinator / Operations Executive / Operations Coordinator", asks for ERP/SAP, MS Office, vendor coordination, KPI monitoring, 0-3 years experience.
   Praghya is a STRONG match. Lead with operational metrics (97% inventory accuracy, 250+ stores, 15% dispatch reduction, 20% turnaround reduction). Use SC vocabulary heavily.

B. supply_chain_analyst_data_heavy
   Signals: title contains "Supply Chain Analyst" / "Data Analyst" AND the JD lists SQL, Python, Power BI, Tableau, dashboards, data modelling, advanced analytics as MUST-HAVE (not nice-to-have).
   Praghya is a WEAK match. Do NOT pretend she has these. Open the summary with operational metrics and analytical foundation from her BBA, NOT with data-tool fluency. Add a `jd_fit_warning` field (see output schema) recommending coordinator-level alternatives. Still produce a usable resume so she can still apply if she wants.

C. graduate_programme_supply_chain
   Signals: "Graduate Programme", "Graduate Scheme", "rotational", "2-year programme" PLUS a supply chain / operations / logistics focus. Examples: Tesco Supply Chain Graduate Scheme, Kerry Europe SC & Customer Care Grad Programme, Infineon Dublin SC Graduate Programme.
   Praghya is a STRONG match. Lead with the MSc + BBA Silver Medal + Oracle SCM cert, then frame Reliance experience as "real operational depth that distinguishes me from typical graduate applicants". Use slightly less heavy SC jargon than archetype A, since these programmes recruit broadly.

D. graduate_programme_business_general
   Signals: "Graduate Programme" / "Business & Commercial Graduate" / "Management Graduate" / consultancy graduate scheme. Examples: Bank of Ireland Business & Commercial Graduate, Big 4 (PwC, Deloitte, EY, KPMG, BDO, Grant Thornton) graduate programmes, Diageo Commercial & Marketing Graduate, Accenture business & technology integration. Often "any degree" / "all disciplines".
   Praghya is a STRONG match. Lead with the Strategy MSc from DCU. Use BUSINESS-STRATEGIC language, not heavy SC jargon. Frame Reliance as "two years of operational and people-leadership experience at one of India's largest retailers" rather than warehouse-detail. Highlight the BBA Silver Medal + 6-month promotion as evidence of consistent excellence.

E. operations_executive_or_inventory
   Signals: "Operations Executive / Operations Coordinator / Inventory Controller / Inventory Coordinator / Materials Coordinator / Warehouse Coordinator". Often retail, FMCG, hospitality, manufacturing.
   Praghya is a STRONG match. Lean into Reliance directly, since these JDs love her exact experience. Use cycle counting, stock reconciliation, GRN, dispatch metrics.

F. pharma_supply_chain
   Signals: pharma / medtech / life sciences company (Pfizer, Abbott, J&J, GSK, Viatris, ABBVIE, MSD, Gilead, West Pharma, Catalyx, Bristol Myers, Bimeda, Uniphar) PLUS supply chain / coordinator / specialist focus. May mention GDP, GMP, regulated environments.
   Praghya is a MODERATE match. Frame her retail experience as "transferable to GDP-regulated supply chain environments". Do NOT claim hands-on GDP/GMP. Emphasise process discipline, documentation rigor, SOP development. Acknowledge in the summary that she would bring strong operational fundamentals adaptable to a regulated context.

G. unknown_stretch
   Signals: domain Praghya has no experience in (construction, data centres, finance trading, niche tech) OR seniority well above her level despite filters.
   Soften the framing. Lead with transferable skills (coordination, process improvement, stakeholder management). Add a `jd_fit_warning`. Still produce the resume so she can decide whether to apply.

═══ STEP 2: TAILOR USING THE ARCHETYPE ═══

Once you've picked an archetype, follow the framing rules above when writing the summary, bullets, and skills.

═══ OUTPUT, RETURN A JSON OBJECT WITH THIS EXACT SHAPE ═══

{
  "jd_archetype": "<one of: supply_chain_coordinator | supply_chain_analyst_data_heavy | graduate_programme_supply_chain | graduate_programme_business_general | operations_executive_or_inventory | pharma_supply_chain | unknown_stretch>",

  "jd_fit_warning": "<empty string if Praghya is a strong match. Otherwise 1-2 short sentences explaining the mismatch and what alternative role at this company she could pursue. Examples: 'This JD requires SQL/Python as must-have, which Praghya does not have. Consider applying to the Supply Chain Coordinator or Operations Executive roles at this company instead.' OR 'This is a construction-sector role outside Praghya's retail/FMCG background; her process and coordination skills are transferable but the domain fit is weak.'>",

  "candidate_title": "<job-title-style line directly under the name. Mirror the JD's role title where possible. Examples: 'Supply Chain Coordinator', 'Logistics Coordinator', 'Operations Executive', 'Graduate Supply Chain', 'Business & Commercial Graduate'. Default to 'Supply Chain & Operations Professional' if unclear. Drop 'Senior'/'Lead'/'Principal' modifiers from the JD title.>",

  "summary": "<EXACTLY 4-5 sentences. Natural flow, no choppy listing. NO em dashes. NO pronouns (no 'I', no 'she', no 'Praghya', no 'her', no 'my'). Drop the subject entirely. The framing depends on the archetype:

  For archetypes A (coordinator), E (operations/inventory), F (pharma):
    Sentence 1: Role identity matched to JD + '2+ years' or 'over 2 years' of experience. Pronoun-free.
       GOOD: 'Supply chain professional with over 2 years of distribution centre experience...'
    Sentence 2: Reliance Retail fulfilment ops for 250+ stores, with one concrete metric (97% inventory accuracy OR 15% dispatch reduction OR 20% turnaround improvement). Pick the one most relevant to the JD.
       GOOD: 'Managed fulfilment for 250+ stores at Reliance Retail, maintaining 97% inventory accuracy.'
    Sentence 3: Education: MSc Management (Strategy), NFQ Level 9 from DCU, plus BBA in Logistics & SCM (9.2/10 GPA, Silver Medallist).
       GOOD: 'Holds an MSc in Management (Strategy), NFQ Level 9, from Dublin City University and a BBA in Logistics & SCM (9.2/10 GPA, Silver Medallist).'
    Sentence 4: Oracle Fusion Cloud SCM certification + which capabilities/tools from the JD she brings.
       GOOD: 'Oracle Fusion Cloud SCM certified, with daily working knowledge of SAP for inventory and KPI monitoring.'
    Sentence 5: Stamp 1G work authorisation note OR a connection to the target company. Pick whichever lands best for this JD.
       GOOD: 'Eligible for full-time employment in Ireland under Stamp 1G, with no employment permit required.'

  For archetypes C (graduate SC programme) and D (graduate business programme):
    Sentence 1: Recent MSc graduate framing. Pronoun-free.
       GOOD: 'Recent MSc graduate of Management (Strategy), NFQ Level 9 from Dublin City University, with a BBA in Logistics & SCM (9.2/10 GPA, Silver Medallist, Top 2 of class).'
    Sentence 2: Verb-led description of Reliance experience. Frame as a differentiator from typical graduate applicants.
       GOOD: 'Brings real operational depth from over 2 years at Reliance Retail (Quick Supply Chain Division) managing fulfilment for 250+ stores.'
       BAD:  'She brings real operational depth...'
    Sentence 3: 6-month promotion as evidence of fast learning. Pronoun-free.
       GOOD: 'Promoted from Graduate Trainee to Senior Graduate within 6 months, demonstrating fast learning and consistent performance.'
       BAD:  'Praghya was promoted...'
    Sentence 4: Why this programme. Connect ambitions to what the programme offers (rotation, breadth, leadership development). Mention the company by name. Verb-led, no pronoun.
       GOOD: 'Drawn to the [Programme Name] for its three-rotation structure...'
    Sentence 5: Stamp 1G eligibility. Important for graduate programmes that often state 'continuous right to work required'. Pronoun-free.
       GOOD: 'Eligible for full-time employment in Ireland under Stamp 1G, with no employment permit required.'

  For archetype B (analyst data-heavy):
    Sentence 1: Frame as 'Operational supply chain professional with strong analytical foundation from a 9.2/10 BBA in SCM and an MSc in Management (Strategy), NFQ Level 9 from DCU'. Lead with thinking ability, not data tools. Pronoun-free.
    Sentence 2: Reliance metrics (97% inventory accuracy, 250+ stores). Verb-led, no pronoun.
       GOOD: 'Managed fulfilment for 250+ stores at Reliance Retail, maintaining 97% inventory accuracy across ~5,000 SKUs.'
    Sentence 3: Acknowledge the analytical aspect honestly: 'Comfortable using Excel and SAP for daily KPI monitoring, with growing exposure to supply chain analytics through ongoing learning.' Pronoun-free.
    Sentence 4: Oracle SCM cert + Stamp 1G note. Pronoun-free.
       GOOD: 'Oracle Fusion Cloud SCM certified. Eligible for full-time employment in Ireland under Stamp 1G, with no employment permit required.'

  For archetype G (unknown stretch):
    Sentence 1: 'Operations and coordination professional with over 2 years...' Pronoun-free.
    Sentence 2: Briefest Reliance metric. Verb-led.
    Sentence 3: Education + Oracle cert. Verb-led.
    Sentence 4: 'Open to applying transferable supply chain, coordination, and process-improvement skills to new sectors.' + Stamp 1G note.

  Avoid robotic phrasing across all archetypes. No 'leveraging', 'utilizing', 'spearheading', 'passionate about', 'committed to excellence', 'seamless', 'innovative'. Write like a confident human.>",

  "skills_additions": {
    "<existing category name>": ["<JD-relevant skill she actually has>", "..."],
    "...": []
  },

  "experience_bullets": [
    {
      "company": "Reliance Retail (Quick Supply Chain Division)",
      "responsibilities": ["<5 rewritten bullets>"],
      "achievements": ["<2 rewritten achievement bullets>"]
    },
    {
      "company": "Om Logistics",
      "responsibilities": ["<4 rewritten bullets>"],
      "achievements": []
    },
    {
      "company": "Shubh Consultants & Technocrats LLP",
      "responsibilities": ["<3 rewritten bullets>"],
      "achievements": []
    }
  ],

  "target_company": "<company name from JD, or 'Company' if not stated>"
}

═══ RULES FOR EACH FIELD ═══

SKILLS_ADDITIONS:
- The keys MUST be one of: "Supply Chain Management", "Operations & Process Improvement", "ERP & Software", "Leadership & Coordination", "Analytical & Research".
- Add ONLY skills that appear in the JD AND that Praghya genuinely has (or has equivalent demonstrated experience).
- Examples of valid additions: "Demand Planning Awareness", "3PL Coordination", "ERP Reporting", "Stakeholder Reporting", "Supplier Onboarding", "Goods-In Documentation", "Service Level Reporting", "Continuous Improvement".
- Do NOT add: Python, SQL, Power BI, Tableau, AWS, Azure, GCP, advanced analytics, machine learning, VBA, S&OP (hands-on), GDP/GMP (hands-on). Even if the JD asks for them.
- If a category has no relevant additions, return an empty list for that category.
- 0-3 additions per category. Quality over quantity.

EXPERIENCE_BULLETS:
- Rewrite each bullet so its WORDING aligns with the JD's archetype, but every concrete claim must come from the BASE responsibilities/achievements provided.
- Every metric stays IDENTICAL: 250+ stores, ~5,000 SKUs, 97% accuracy, ~15%, ~20%, 20+ staff, ~10%, ~200+ weekly consignments, 6 months. Never change a number.
- Reliance Retail: keep 5 responsibility bullets and 2 achievement bullets.
- Om Logistics: keep 4 responsibility bullets, 0 achievement bullets.
- Shubh Consultants: keep 3 responsibility bullets, 0 achievement bullets.
- Each bullet starts with a strong past-tense verb: managed, oversaw, coordinated, tracked, maintained, reduced, improved, trained, processed, verified, supported, streamlined.
- Do NOT introduce new tools, sectors, or claims not in the base bullets. You may RE-LABEL an existing claim using a JD-aligned synonym (for example, 'warehouse operations' becomes 'distribution centre operations' if JD uses that language; 'fulfilment centre' and 'distribution centre' interchange freely; 'cycle counts' becomes 'cycle counting'; 'goods receipt' becomes 'GRN processing') but the underlying fact must match.

INDUSTRY BRIDGE, when JD industry differs from retail:
- Pharma / medtech / life sciences (Catalyx, Abbott, J&J, etc.): use phrasing like 'distribution centre operations and stock reconciliation discipline transferable to regulated supply chain environments'. Do NOT claim GDP/GMP hands-on experience.
- FMCG / retail / food (Tesco, Kerry, Diageo, Pernod Ricard, Aramark, Ornua): lead with Reliance directly, since it's the same domain. Use freely.
- Logistics / 3PL / freight forwarding (DFDS, Expeditors, Constellation): emphasise Om Logistics consignment tracking + Reliance dispatch coordination + 3PL liaison.
- Tech / data centre / construction: soften the framing, mention transferable skills only. Set jd_archetype to 'unknown_stretch'.
- Financial services / consulting / Big 4 / banking grad programme: business-strategic language, lead with MSc Strategy. Frame Reliance experience as 'operational case-study experience: 2 years inside one of India's largest retail supply chains'.

- NO em dashes inside bullets. Use commas, periods, colons, or parentheses.

SUMMARY:
- 4-5 sentences. Count them.
- Mention the target company by name if it appears clearly in the JD.
- Include the Stamp 1G work authorisation note unless it would awkwardly displace something more important.
- One concrete metric from her real experience.
- Confident but not boastful.
- NO em dashes.

CANDIDATE_TITLE:
- Match the JD's role title verbatim where reasonable (max 8 words).
- Drop seniority modifiers ('Senior', 'Lead') if present in the JD title, since Praghya is entry-to-junior level.

JD_ARCHETYPE & JD_FIT_WARNING:
- Always populate jd_archetype with one of the seven values listed.
- jd_fit_warning is empty string '' when archetype is A, C, D, E, F (good fit). Populate it for B and G with a concise honest assessment + alternative role suggestion.
- NO em dashes in the warning text either.

═══ FINAL CHECK BEFORE OUTPUT ═══
Scan your entire output for the em dash character (—). If you find any, remove or replace them. This applies to summary, bullets, candidate_title, jd_fit_warning, and any other field.

═══ OUTPUT ═══
Return ONLY the JSON object. No prose, no markdown fences, no explanation.
"""


COVER_LETTER_PROMPT = """You are Praghya Prakhar writing a cover letter for the role described below.
Write in first person. Sound like a real human, not a corporate template.

═══ HARD RULE: NO EM DASHES ANYWHERE IN THE LETTER ═══

This is non-negotiable. Em dashes (—, U+2014) are a leading recruiter signal of AI-generated writing in 2026. Do NOT use them anywhere in the letter body. Use commas, periods, colons, or sentence breaks instead. Hyphens inside compound words (e.g. 'cross-functional', 'on-time-in-full', 'first-hand') are fine. En dashes inside date ranges are fine. Em dashes are NOT.

Before you finalise, scan your letter for the em dash character (—) and remove every one.

═══ CONTEXT ═══
Praghya is an entry-to-junior level Supply Chain & Operations professional based in Dublin.
- 2+ years at Reliance Retail (Quick Supply Chain Division), Delhi: distribution centre operations for 250+ stores, 97% inventory accuracy across ~5,000 SKUs, ~15% reduction in dispatch delays, ~20% reduction in goods-in turnaround time, GRN processing, 3PL coordination, trained 20+ warehouse staff.
- Promoted from Graduate Trainee to Senior Graduate within 6 months at Reliance.
- Internships at Om Logistics (consignment tracking, ~10% clearance turnaround improvement, ~200+ weekly consignments verified, PO verification) and Shubh Consultants (project coordination).
- MSc in Management (Strategy), NFQ Level 9, from Dublin City University, Grade 2:1.
- BBA in Logistics & Supply Chain Management from Galgotias University, 9.2/10 GPA, Silver Medallist (Top 2 of class).
- Oracle Fusion Cloud SCM certified (Nov 2025).
- Eligible for full-time employment in Ireland under Stamp 1G (Third Level Graduate Programme), with no employment permit required.

═══ HARD RULES ═══

DO NOT mention skills she does not have:
- No Python, SQL, R, JavaScript, Java
- No Power BI, Tableau, Looker, Qlik
- No advanced Excel (no VBA, no macros)
- No AWS, Azure, GCP, databases, cloud platforms
- No machine learning, data science, advanced analytics
- No Six Sigma certification or PMP certification (she is currently learning, NOT certified)
- No hands-on S&OP experience, no hands-on GDP/GMP regulatory experience

DO NOT inflate experience:
- She has 2+ years of experience. Not 3, not 5.
- Her experience is in retail fulfilment / warehouse ops / logistics. If the JD is in pharma, construction, or another sector, FRAME her retail experience as transferable. Never claim sector-specific experience she does not have.

═══ DETECT THE JD TYPE FIRST ═══

Identify which category the JD falls into. It changes the opening hook and war story:
A. Direct-hire SC/operations role (coordinator, specialist, executive, inventory): use operational war story
B. Graduate programme, supply chain (Tesco, Kerry, Infineon Supply Chain Grad): lead with education + curiosity for the rotation
C. Graduate programme, business/commercial (Bank of Ireland, Big 4, Diageo, Accenture): lead with strategy MSc + transferable Reliance ops experience
D. Pharma / regulated supply chain: frame retail experience as transferable to GDP-regulated environments
E. Stretch role (different domain): soften, lean on transferable skills

═══ BANNED PHRASES ═══
Do not use any of these (they make the letter sound AI-generated or template-y):
- "I am writing to express my interest"
- "I am excited to apply"
- "Please find my resume attached"
- "I believe I am a perfect fit"
- "passionate about", "driven by", "committed to excellence"
- "leveraging", "utilizing", "harnessing"
- "showcasing", "highlighting", "demonstrating"
- "testament to", "underscores", "pivotal", "tapestry", "realm"
- "seamless", "innovative", "groundbreaking", "cutting-edge"
- "at the forefront of", "at the intersection of"
- Three-adjective chains ("scalable, reliable, and efficient")
- ANY em dash (—)

═══ STRUCTURE ═══
4 short paragraphs, plain text, no markdown, no headers, no bold.

Paragraph 1 (Hook, 2-3 sentences):
Open by referring to a SPECIFIC operational challenge or focus from the JD (not the company in general). Show you actually read what they wrote.

Examples for direct-hire SC/operations roles (type A):
- "Keeping inventory accuracy above 95% across hundreds of SKUs is harder than most people think. It depends entirely on the cycle counting discipline behind the scenes."
- "Coordinating 3PL deliveries against rolling forecast changes is exactly the kind of problem I worked on every day at Reliance Retail."

Examples for graduate programmes (types B and C):
- "What drew me to the [programme name] is the structure of three rotations across operations, analytics, and projects. That breadth is exactly what I want in my first Irish role."
- "A graduate programme that pairs structured rotations with real responsibility from day one is rare, and the [Bank of Ireland Business & Commercial / Tesco Supply Chain / Diageo Commercial] programme stands out for that reason."

Mention the role title from the JD and the company name in this paragraph.

Paragraph 2 (War story, 3-4 sentences):
Pick the BEST matching war story from her real experience, based on the JD's focus:

- If the JD is heavy on inventory/SKU management: use the 97% inventory accuracy + 5,000 SKU cycle counting story.
- If the JD is about efficiency/process improvement: use the 20% goods-in turnaround OR 15% dispatch delay reduction story.
- If the JD is about people leadership / training / onboarding: use the "trained 20+ warehouse staff" + "Graduate Trainee to Senior Graduate in 6 months" story.
- If the JD is about vendor / 3PL / supplier coordination: use the Om Logistics 200+ weekly consignments + 10% clearance turnaround story, paired with Reliance 3PL coordination.
- If the JD is a graduate SUPPLY CHAIN programme (type B): describe the Reliance promotion in 6 months as evidence of how quickly she ramps up, and connect to one operational metric (97% accuracy or 20% turnaround).
- If the JD is a graduate BUSINESS programme (type C, Big 4/Bank of Ireland/Diageo): frame Reliance more strategically: 'two years inside one of India's largest retailers gave me first-hand exposure to how operational decisions cascade across 250+ store fronts'. Connect to the BBA Silver Medal as evidence of academic excellence.

Use exact metrics. Never round or change numbers.

Paragraph 3 (Education + Certification, 2 sentences):
Mention the MSc in Management (Strategy), NFQ Level 9 from DCU, and the Oracle Fusion Cloud SCM certification (or BBA Silver Medal, depending on what fits). Connect the certification or coursework to a tool, process, or theme the JD asks for.

Paragraph 4 (Close, 2-3 sentences):
Brief, confident close. Mention Stamp 1G eligibility ('I am eligible to work full-time in Ireland under Stamp 1G, with no employment permit required.'). This is a real differentiator and graduate programmes often ask about it. Express interest in discussing the role. End with "Thank you," on its own line, then "Praghya Prakhar" on the next line.

═══ STYLE ═══
- Vary sentence length. Mix short with longer.
- NO em dashes. Use commas, periods, colons, or rephrase.
- Use plain verbs: managed, coordinated, tracked, maintained, reduced, improved, trained, processed, verified.
- She is entry-level. Confident, not arrogant. Eager to learn, not desperate.
- Length: 240-340 words total in the letter body (excluding "Thank you, Praghya Prakhar" sign-off).

═══ FINAL CHECK ═══
Before you output, scan your letter once more for the em dash character (—) and remove every instance.

═══ OUTPUT ═══
Return ONLY the letter body as plain text. No "Dear Hiring Manager" greeting (the renderer adds it). No subject line. No address blocks. No markdown. No bold. No code fences.
Start directly with paragraph 1.
End with "Thank you," on its own line and "Praghya Prakhar" on the next line.
"""


# ═══════════════════════════════════════════════════════════════════
# 5. SCHEMA
# ═══════════════════════════════════════════════════════════════════

class ExperienceBullets(BaseModel):
    company: str
    responsibilities: List[str]
    achievements: List[str] = Field(default_factory=list)


class SkillsAdditions(BaseModel):
    supply_chain_management: List[str] = Field(default_factory=list, alias="Supply Chain Management")
    operations_process: List[str] = Field(default_factory=list, alias="Operations & Process Improvement")
    erp_software: List[str] = Field(default_factory=list, alias="ERP & Software")
    leadership: List[str] = Field(default_factory=list, alias="Leadership & Coordination")
    analytical: List[str] = Field(default_factory=list, alias="Analytical & Research")

    class Config:
        populate_by_name = True


class TailoredOutput(BaseModel):
    candidate_title: str
    summary: str
    skills_additions: dict  # {category_name: [skills]}
    experience_bullets: List[ExperienceBullets]
    target_company: str = "Company"
    jd_archetype: str = ""
    jd_fit_warning: str = ""


# ═══════════════════════════════════════════════════════════════════
# 6. CALL GEMINI
# ═══════════════════════════════════════════════════════════════════

def call_gemini(api_key: str, jd_text: str, resume_text: str = "") -> dict:
    """Single API call. JSON mode. Returns parsed dict or {'error': ...}.
    resume_text is Praghya's base resume — passed through so the model
    sees what she sees in the UI box and respects any edits she made there.
    Falls back to the constants-derived default if not supplied."""
    if not api_key:
        return {"error": "Missing GOOGLE_API_KEY."}
    if not jd_text or not jd_text.strip():
        return {"error": "Job description is empty."}

    if not resume_text or not resume_text.strip():
        resume_text = PRAGHYA_BASE_RESUME

    client = genai.Client(api_key=api_key)
    prompt = (
        f"{ASTRA_PROMPT}\n\n"
        f"═══ BASE RESUME ═══\n{resume_text}\n\n"
        f"═══ JOB DESCRIPTION ═══\n{jd_text}"
    )

    try:
        response = client.models.generate_content(
            model=GENERATION_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.4,  # low for reliability
            ),
        )
        text = response.text.strip()
        # Strip code fences if model wrapped them despite JSON mode
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        return json.loads(text)
    except json.JSONDecodeError as e:
        return {"error": f"Could not parse JSON from model: {e}"}
    except Exception as e:
        return {"error": f"Generation error: {e}"}


def generate_cover_letter(api_key: str, resume_data: dict, jd_text: str) -> str:
    """Single API call. Plain text mode. Returns letter body or error message starting with 'ERROR:'."""
    if not api_key:
        return "ERROR: Missing GOOGLE_API_KEY."
    if not jd_text or not jd_text.strip():
        return "ERROR: Job description is empty."

    client = genai.Client(api_key=api_key)

    # Pass the tailored resume context so the cover letter aligns with what the JD-tailored CV says.
    resume_context = (
        f"Tailored role title: {resume_data.get('candidate_title', '')}\n"
        f"Target company: {resume_data.get('target_company', '')}\n"
        f"Tailored summary: {resume_data.get('summary', '')}"
    )

    prompt = (
        f"{COVER_LETTER_PROMPT}\n\n"
        f"═══ TAILORED RESUME CONTEXT ═══\n{resume_context}\n\n"
        f"═══ JOB DESCRIPTION ═══\n{jd_text}"
    )

    try:
        response = client.models.generate_content(
            model=GENERATION_MODEL,
            contents=prompt,
            config=types.GenerateContentConfig(temperature=0.6),
        )
        text = (response.text or "").strip()
        # Strip code fences if any
        if text.startswith("```"):
            text = re.sub(r"^```(?:\w+)?\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
        # Strip any "Dear ..." greeting if the model added one despite instructions
        text = re.sub(r"^dear\s+[^\n]+\n+", "", text, flags=re.IGNORECASE)
        # Final scrub of banned skills (safety net)
        text = scrub_banned_from_text(text)
        return text
    except Exception as e:
        return f"ERROR: Cover letter generation failed: {e}"


# ═══════════════════════════════════════════════════════════════════
# 7. ASSEMBLE: merge model output with base resume into final structure
# ═══════════════════════════════════════════════════════════════════

def is_banned(skill: str) -> bool:
    s = skill.lower().strip()
    return any(b in s for b in BANNED_SKILLS)


def strip_em_dashes(text: str) -> str:
    """Replace em dashes (U+2014) with cleaner alternatives. Em dashes are a
    leading recruiter signal of AI-generated writing in 2026; we want zero in
    rendered output. This is the runtime safety net in case the model slips
    one through despite the prompt instructions.

    Replacement strategy:
    - " — " (em dash with spaces) -> ", " (comma + space) which reads naturally in 95% of cases
    - "word—word" (em dash without spaces) -> "word, word"

    We deliberately do NOT touch hyphens (-, U+002D) or en dashes (–, U+2013).
    Hyphens are correct in compound words (cross-functional). En dashes are
    correct in date ranges (Aug 2022 – Dec 2024). Neither is an AI signal.
    """
    if not text:
        return text
    # Replace em dash + surrounding whitespace with comma + space
    cleaned = re.sub(r"\s*—\s*", ", ", text)
    # Tidy up the rare case where the em dash followed punctuation, leaving ", ,"
    cleaned = re.sub(r",\s*,", ",", cleaned)
    return cleaned


def strip_summary_pronouns(text: str) -> str:
    """Remove first-person/third-person pronouns from the start of summary
    sentences. Resume summaries follow the Irish/UK convention of pronoun-free,
    implied-subject prose ("Operations professional with...", not "I am an
    operations professional..." or "She is an operations professional...").

    This is the runtime safety net. The prompt tells the model to write
    pronoun-free, but if a stray "She brings..." or "I bring..." or "Praghya
    is..." slips through, we catch it before rendering.

    Applied ONLY to the summary string. Cover letters use first-person
    ("I", "my") which is normal letter convention and must NOT be stripped.
    """
    if not text:
        return text

    # Sentence-initial pronoun patterns to remove. Order matters: longer phrases first
    # so 'She is a' doesn't get partially matched by 'She '.
    # The patterns below match at the start of the string OR after a period+space,
    # which is the start of a sentence.
    patterns = [
        # "She is a/an X" / "Praghya is a/an X" / "I am a/an X" -> "X"
        (r"(^|(?<=\.\s))(I am|I'm|She is|She's|He is|He's|Praghya is|Praghya was)\s+(a|an)\s+", r"\1"),
        # "She is X" / "Praghya is X" / "I am X" -> "X"
        (r"(^|(?<=\.\s))(I am|I'm|She is|She's|He is|He's|Praghya is|Praghya was)\s+", r"\1"),
        # "She brings/has/holds/manages X" -> "Brings/Has/Holds X" (verb stays, capitalised)
        (r"(^|(?<=\.\s))(She|He)\s+(brings|has|holds|manages|managed|coordinates|coordinated|oversees|oversaw|maintains|maintained|trains|trained|reduced|improved|streamlined|verified|tracked|delivered|leads|led|supports|supported|drives|drove)\b",
         lambda m: m.group(1) + m.group(3).capitalize()),
        # "I bring X" / "I have X" / etc. -> "Bring X" / "Have X" / "Hold X"
        (r"(^|(?<=\.\s))I\s+(bring|have|hold|manage|managed|coordinate|coordinated|oversee|oversaw|maintain|maintained|train|trained|reduce|reduced|improve|improved|streamline|streamlined|verify|verified|track|tracked|deliver|delivered|lead|led|support|supported|drive|drove)\b",
         lambda m: m.group(1) + m.group(2).capitalize()),
        # "Praghya brings/managed/etc X" (verb retained) -> "Brings/Managed X"
        (r"(^|(?<=\.\s))Praghya\s+(managed|coordinated|oversaw|maintained|trained|reduced|improved|streamlined|verified|tracked|delivered|led|supported|drove|brings|has|holds|manages|coordinates|oversees|maintains|trains|reduces|improves|streamlines|verifies|tracks|delivers|leads|supports|drives)\b",
         lambda m: m.group(1) + m.group(2).capitalize()),
        # Possessives: "her experience" / "my MSc" -> "experience" / "MSc"
        (r"\b(?:her|My|my)\s+(MSc|BBA|experience|background|certification)\b", r"\1"),
    ]

    cleaned = text
    for pattern, replacement in patterns:
        cleaned = re.sub(pattern, replacement, cleaned)

    # First letter of result should be uppercase (the first sentence)
    if cleaned and cleaned[0].islower():
        cleaned = cleaned[0].upper() + cleaned[1:]

    # Tidy: collapse double spaces that may have been introduced
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    return cleaned


def scrub_banned_from_text(text: str) -> str:
    """Strip any banned skill mentions and em dashes from a free-text string
    (summary, bullets, cover letter). Two-stage scrub:
    1. Remove banned skills (Python, SQL, Power BI, etc.) - safety net for
       cases where the model added a banned skill despite the prompt
    2. Strip all em dashes - safety net so rendered output never has the
       AI-generated-text signal
    """
    cleaned = text
    for b in BANNED_SKILLS:
        # Remove standalone occurrences with surrounding context cleanup
        pattern = re.compile(r"\b" + re.escape(b) + r"\b[, ]*", re.IGNORECASE)
        cleaned = pattern.sub("", cleaned)
    # Tidy up double commas, double spaces, trailing punctuation
    cleaned = re.sub(r",\s*,", ",", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned)
    cleaned = re.sub(r"\s+([,.])", r"\1", cleaned)
    # Strip em dashes - applied LAST so the previous tidy-ups don't reintroduce any
    cleaned = strip_em_dashes(cleaned)
    return cleaned.strip()


def _normalize_skill(s: str) -> str:
    """Normalise a skill for dedup comparison. Strips parens, punctuation,
    and lowercases. So 'SAP (Inventory & SCM Modules)' and 'SAP Inventory
    Modules' both contain the normalised core 'sap inventory'.
    """
    if not s:
        return ""
    # Lowercase, strip parens content keeping the words inside, normalise spaces
    n = s.lower()
    # Remove parens but keep contents: "SAP (Inventory & SCM Modules)" -> "sap inventory & scm modules"
    n = re.sub(r"[()]", " ", n)
    # Drop punctuation that doesn't carry meaning for matching
    n = re.sub(r"[^a-z0-9& ]", " ", n)
    # Collapse whitespace
    n = re.sub(r"\s+", " ", n).strip()
    return n


def _is_duplicate_skill(candidate: str, existing_normalized: set) -> bool:
    """Return True if `candidate` is a duplicate of a skill already in
    `existing_normalized`. Handles three cases:
    1. Exact match after normalisation: 'SAP' vs 'sap'
    2. Candidate is a substring of an existing skill: 'SAP' vs 'sap inventory & scm modules'
    3. An existing skill is a substring of candidate: 'sap' (existing) vs 'SAP Inventory Modules' (candidate)

    The substring check uses word-boundary token containment to avoid false
    positives like 'Management' matching 'Vendor Management' (those ARE distinct
    skills; we use full-token containment, not substring containment).
    """
    cand_norm = _normalize_skill(candidate)
    if not cand_norm:
        return True  # empty / pure-punctuation skills are duplicates of nothing
    if cand_norm in existing_normalized:
        return True

    cand_tokens = set(cand_norm.split())
    for existing_norm in existing_normalized:
        existing_tokens = set(existing_norm.split())
        if not existing_tokens or not cand_tokens:
            continue
        # Case A: candidate is a "smaller version" of an existing multi-word skill
        # ("sap" is a subset of {"sap", "inventory", "scm", "modules"})
        if cand_tokens.issubset(existing_tokens):
            return True
        # Case B: existing is a "smaller version" of candidate
        # (existing "sap" makes new "sap inventory modules" redundant)
        if existing_tokens.issubset(cand_tokens):
            return True
    return False


def merge_skills(additions: dict) -> dict:
    """Combine BASE_SKILLS with model additions. Drops banned terms and
    cross-category duplicates so the rendered skills section is clean.

    Dedup rules:
    - A model-added skill is rejected if it already exists in ANY category
      (base or other addition) by the _is_duplicate_skill check
    - A model-added skill is rejected if it's banned (Python, SQL, etc.)
    - A model-added skill is rejected if it's a token-subset or token-superset
      of any existing skill in any category
    """
    final = {}
    # Build the GLOBAL set of normalised skills across all base categories first
    # so cross-category dedup works.
    seen_normalized = set()
    for category, base_list in BASE_SKILLS.items():
        for skill in base_list:
            seen_normalized.add(_normalize_skill(skill))

    # Start each output category with its base skills (always preserved as-is)
    for category, base_list in BASE_SKILLS.items():
        final[category] = list(base_list)

    # Now pass through model additions, deduping against the global set
    if not isinstance(additions, dict):
        additions = {}

    for category, added_list in additions.items():
        if category not in final:
            continue  # ignore unknown categories
        if not isinstance(added_list, list):
            continue
        for skill in added_list:
            if not isinstance(skill, str):
                continue
            s = skill.strip()
            if not s:
                continue
            if is_banned(s):
                continue
            if _is_duplicate_skill(s, seen_normalized):
                continue
            # Accept the addition
            final[category].append(s)
            # Update global set so a subsequent category can't re-add it
            seen_normalized.add(_normalize_skill(s))
    return final


def merge_experience(model_bullets: list) -> list:
    """For each base role, replace bullets with the model's tailored versions
    (after scrubbing). If the model dropped a role or returned the wrong count,
    fall back to base bullets so the resume stays complete."""
    by_company = {}
    if isinstance(model_bullets, list):
        for item in model_bullets:
            if isinstance(item, dict) and item.get("company"):
                by_company[item["company"]] = item

    final = []
    for base_role in BASE_EXPERIENCE:
        company = base_role["company"]
        tailored = by_company.get(company, {})
        tailored_resps = tailored.get("responsibilities") or []
        tailored_achs = tailored.get("achievements") or []

        # Scrub banned terms; drop empty
        clean_resps = [scrub_banned_from_text(r) for r in tailored_resps if isinstance(r, str)]
        clean_resps = [r for r in clean_resps if r.strip()]

        clean_achs = [scrub_banned_from_text(a) for a in tailored_achs if isinstance(a, str)]
        clean_achs = [a for a in clean_achs if a.strip()]

        # Fallback: if tailoring lost too many bullets, use base
        expected_resp_count = len(base_role["responsibilities"])
        expected_ach_count = len(base_role["achievements"])

        if len(clean_resps) < max(1, expected_resp_count - 1):
            clean_resps = list(base_role["responsibilities"])
        if expected_ach_count > 0 and len(clean_achs) < expected_ach_count:
            clean_achs = list(base_role["achievements"])

        final.append({
            "role_title": base_role["role_title"],
            "company": base_role["company"],
            "location": base_role["location"],
            "dates": base_role["dates"],
            "responsibilities": clean_resps,
            "achievements": clean_achs,
        })
    return final


def assemble_resume(model_output: dict) -> dict:
    """Combine model output + base facts into a fully-populated resume dict."""
    summary = scrub_banned_from_text(model_output.get("summary", "") or "")
    summary = strip_summary_pronouns(summary)
    if not summary or len(summary.split()) < 20:
        # Fallback summary if model failed
        summary = (
            "Operations and supply chain professional with over 2 years of hands-on "
            "experience in warehouse and distribution centre operations, inventory "
            "control, and logistics coordination. Held a Senior Graduate role at "
            "Reliance Retail, overseeing fulfilment for 250+ stores and maintaining "
            "97% inventory accuracy across ~5,000 SKUs through systematic cycle "
            "counting and stock reconciliation. Holds an MSc in Management "
            "(Strategy), NFQ Level 9, from Dublin City University and a BBA in "
            "Logistics & Supply Chain Management with a 9.2/10 GPA and Silver "
            "Medal. Oracle Fusion Cloud SCM certified, with daily working "
            "knowledge of SAP and Excel for inventory, KPI monitoring, and "
            "service-level reporting. Eligible for full-time employment in "
            "Ireland under Stamp 1G, with no employment permit required."
        )

    return {
        "candidate_name": CANDIDATE_NAME,
        "candidate_title": strip_em_dashes((model_output.get("candidate_title") or CANDIDATE_TAGLINE).strip()),
        "contact_info": CANDIDATE_CONTACT,
        "summary": summary,
        "skills": merge_skills(model_output.get("skills_additions", {}) or {}),
        "experience": merge_experience(model_output.get("experience_bullets", []) or []),
        "education": list(BASE_EDUCATION),
        "certifications": list(BASE_CERTIFICATIONS),
        "additional_info": list(BASE_ADDITIONAL_INFO),
        "target_company": strip_em_dashes((model_output.get("target_company") or "Company").strip()),
        "jd_archetype": (model_output.get("jd_archetype") or "").strip(),
        "jd_fit_warning": strip_em_dashes((model_output.get("jd_fit_warning") or "").strip()),
    }


# ═══════════════════════════════════════════════════════════════════
# 8. DOCX RENDERER — clean, ATS-friendly, single-column, no tables
# ═══════════════════════════════════════════════════════════════════

def _set_font(run, size, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except Exception:
        pass


def _add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(text), 12, bold=True)


def _add_bullet(doc, text, bold_prefix=None, justify=True):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY if justify else WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_after = Pt(0)
    if bold_prefix:
        _set_font(p.add_run(bold_prefix), 12, bold=True)
    _set_font(p.add_run(text), 12)


def render_docx(data: dict) -> bytes:
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)

    # ─── Header ───
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(data["candidate_name"])
    run.font.all_caps = True
    _set_font(run, 28, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    _set_font(p.add_run(data["candidate_title"]), 14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(data["contact_info"]), 11, bold=True)

    # ─── Professional Profile ───
    _add_section_heading(doc, "Professional Profile")
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    _set_font(p.add_run(data["summary"]), 12)

    # ─── Skills ───
    _add_section_heading(doc, "Key Skills / Tools & Technologies")
    for cat, skill_list in data["skills"].items():
        _add_bullet(doc, ", ".join(skill_list), bold_prefix=f"{cat}: ")

    # ─── Experience ───
    _add_section_heading(doc, "Professional Experience")
    for role in data["experience"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        header = f"{role['role_title']} | {role['company']} | {role['location']} | {role['dates']}"
        _set_font(p.add_run(header), 12, bold=True)

        for resp in role["responsibilities"]:
            _add_bullet(doc, resp)

        if role["achievements"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(0)
            _set_font(p.add_run("Achievements:"), 12, bold=True)
            for ach in role["achievements"]:
                _add_bullet(doc, ach)

    # ─── Education ───
    _add_section_heading(doc, "Education")
    for edu in data["education"]:
        # Degree line (bold)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        _set_font(p.add_run(edu["degree"]), 12, bold=True)

        # Institution / dates / grade
        meta = f"{edu['institution']} | {edu['dates']} | {edu['grade']}"
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _set_font(p.add_run(meta), 12)

        if edu.get("extra"):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.paragraph_format.space_after = Pt(0)
            _set_font(p.add_run(edu["extra"]), 12)

    # ─── Certifications ───
    _add_section_heading(doc, "Certifications")
    for cert in data["certifications"]:
        _add_bullet(doc, cert)

    # ─── Additional Information ───
    _add_section_heading(doc, "Additional Information")
    for line in data["additional_info"]:
        _add_bullet(doc, line)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ═══════════════════════════════════════════════════════════════════
# 9. PDF RENDERER — same content, ATS-friendly, single column
# ═══════════════════════════════════════════════════════════════════

def _esc(t):
    if t is None:
        return ""
    return escape(str(t))


def render_pdf(data: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()

    sn = ParagraphStyle("N", parent=styles["Normal"], fontName="Times-Roman",
                        fontSize=12, leading=14, alignment=TA_JUSTIFY, spaceAfter=0)
    sn_left = ParagraphStyle("NL", parent=styles["Normal"], fontName="Times-Roman",
                             fontSize=12, leading=14, alignment=TA_LEFT, spaceAfter=0)
    sh_name = ParagraphStyle("HN", parent=styles["Normal"], fontName="Times-Bold",
                             fontSize=22, leading=26, alignment=TA_CENTER, spaceAfter=0)
    sh_title = ParagraphStyle("HT", parent=styles["Normal"], fontName="Times-Bold",
                              fontSize=14, leading=16, alignment=TA_CENTER, spaceAfter=0)
    sh_contact = ParagraphStyle("HC", parent=styles["Normal"], fontName="Times-Bold",
                                fontSize=11, leading=13, alignment=TA_CENTER, spaceAfter=6)
    s_sec = ParagraphStyle("Sec", parent=styles["Normal"], fontName="Times-Bold",
                           fontSize=12, leading=14, alignment=TA_LEFT,
                           spaceBefore=10, spaceAfter=2)

    el = []

    # Header
    el.append(Paragraph(_esc(data["candidate_name"]).upper(), sh_name))
    el.append(Paragraph(_esc(data["candidate_title"]), sh_title))
    el.append(Paragraph(_esc(data["contact_info"]), sh_contact))

    # Summary
    el.append(Paragraph("Professional Profile", s_sec))
    el.append(Paragraph(_esc(data["summary"]), sn))

    # Skills
    el.append(Paragraph("Key Skills / Tools &amp; Technologies", s_sec))
    skill_items = []
    for cat, skill_list in data["skills"].items():
        text = f"<b>{_esc(cat)}:</b> {_esc(', '.join(skill_list))}"
        skill_items.append(ListItem(Paragraph(text, sn_left), leftIndent=0))
    if skill_items:
        el.append(ListFlowable(skill_items, bulletType="bullet",
                               start="\u2022", leftIndent=15))

    # Experience
    el.append(Paragraph("Professional Experience", s_sec))
    for role in data["experience"]:
        header = f"{role['role_title']} | {role['company']} | {role['location']} | {role['dates']}"
        el.append(Paragraph(f"<b>{_esc(header)}</b>", sn_left))
        el.append(Spacer(1, 2))
        items = [ListItem(Paragraph(_esc(r), sn), leftIndent=0)
                 for r in role["responsibilities"] if r.strip()]
        if items:
            el.append(ListFlowable(items, bulletType="bullet", start="\u2022", leftIndent=15))
        if role["achievements"]:
            el.append(Paragraph("<b>Achievements:</b>", sn_left))
            ach_items = [ListItem(Paragraph(_esc(a), sn), leftIndent=0)
                         for a in role["achievements"] if a.strip()]
            if ach_items:
                el.append(ListFlowable(ach_items, bulletType="bullet",
                                       start="\u2022", leftIndent=25))
        el.append(Spacer(1, 4))

    # Education
    el.append(Paragraph("Education", s_sec))
    for edu in data["education"]:
        el.append(Paragraph(f"<b>{_esc(edu['degree'])}</b>", sn_left))
        meta = f"{edu['institution']} | {edu['dates']} | {edu['grade']}"
        el.append(Paragraph(_esc(meta), sn_left))
        if edu.get("extra"):
            el.append(Paragraph(_esc(edu["extra"]), sn))
        el.append(Spacer(1, 4))

    # Certifications
    el.append(Paragraph("Certifications", s_sec))
    cert_items = [ListItem(Paragraph(_esc(c), sn_left), leftIndent=0)
                  for c in data["certifications"]]
    if cert_items:
        el.append(ListFlowable(cert_items, bulletType="bullet",
                               start="\u2022", leftIndent=15))

    # Additional
    el.append(Paragraph("Additional Information", s_sec))
    add_items = [ListItem(Paragraph(_esc(a), sn_left), leftIndent=0)
                 for a in data["additional_info"]]
    if add_items:
        el.append(ListFlowable(add_items, bulletType="bullet",
                               start="\u2022", leftIndent=15))

    doc.build(el)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════
# 9b. COVER LETTER DOCX RENDERER
# ═══════════════════════════════════════════════════════════════════

def render_cover_letter_docx(letter_body: str, target_company: str = "") -> bytes:
    """Render a cover letter as a clean DOCX. Body is passed in as plain text;
    the renderer adds the candidate header, date, greeting, and sign-off."""
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1.0)

    def _line(text, bold=False, space_after=0, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            _set_font(p.add_run(text), 11, bold=bold)
        return p

    # ─── Sender header ───
    _line(CANDIDATE_NAME, bold=True, space_after=0)
    # Split contact info: "Dublin, Ireland | +353 ... | email | linkedin"
    for piece in [c.strip() for c in CANDIDATE_CONTACT.split("|")]:
        if piece:
            _line(piece, space_after=0)

    _line("", space_after=10)  # gap

    # ─── Date ───
    today = datetime.date.today().strftime("%d %B %Y")
    _line(today, space_after=14)

    # ─── Greeting ───
    if target_company.strip():
        greeting = f"Dear Hiring Team at {target_company.strip()},"
    else:
        greeting = "Dear Hiring Team,"
    _line(greeting, space_after=10)

    # ─── Body paragraphs ───
    # Strip any sign-off the model may have included so we don't render it twice
    body = letter_body.strip()
    sign_off_pattern = re.compile(
        r"\n+(thank you,?|regards,?|sincerely,?|kind regards,?|best regards,?)\s*\n+praghya prakhar\s*$",
        re.IGNORECASE,
    )
    # Detect sign-off, render it separately for nicer spacing
    sign_off_match = sign_off_pattern.search(body)
    if sign_off_match:
        body_main = body[: sign_off_match.start()].strip()
        sign_word = sign_off_match.group(1).rstrip(",").strip().capitalize()
    else:
        body_main = body
        sign_word = "Thank you"

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", body_main) if p.strip()]
    if not paragraphs:
        # Treat single-line newline-separated text as one paragraph each
        paragraphs = [p.strip() for p in body_main.split("\n") if p.strip()]

    for para in paragraphs:
        # Collapse internal newlines into a single space within a paragraph
        clean_para = re.sub(r"\s*\n\s*", " ", para)
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        _set_font(p.add_run(clean_para), 11)

    # ─── Sign-off ───
    _line("", space_after=6)
    _line(f"{sign_word},", space_after=18)
    _line(CANDIDATE_NAME, space_after=0)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ═══════════════════════════════════════════════════════════════════
# 10. STREAMLIT UI
# ═══════════════════════════════════════════════════════════════════

st.set_page_config(page_title=PAGE_TITLE, layout="wide", page_icon="📦",
                   initial_sidebar_state="expanded")

st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .block-container {padding-top: 1.5rem;}
    div.stButton > button:first-child {border-radius: 6px; font-weight: 600;}
</style>
""", unsafe_allow_html=True)

# Session state
if "tailored" not in st.session_state:
    st.session_state["tailored"] = None
if "saved_jd" not in st.session_state:
    st.session_state["saved_jd"] = ""
if "saved_base" not in st.session_state:
    st.session_state["saved_base"] = PRAGHYA_BASE_RESUME
if "cover_letter" not in st.session_state:
    st.session_state["cover_letter"] = None

# Sidebar
with st.sidebar:
    st.header("⚙️ Configuration")
    if GOOGLE_API_KEY:
        st.success("API key configured")
        api_key = GOOGLE_API_KEY
    else:
        st.warning("Add GOOGLE_API_KEY to Streamlit secrets, or paste it below.")
        api_key = st.text_input("Google API Key", type="password")

    st.divider()
    st.markdown("**Target roles:**")
    st.caption(
        "**Track A, Direct hire:** Supply Chain Coordinator · Operations Executive · "
        "Inventory Coordinator · Logistics Coordinator\n\n"
        "**Track B, Graduate programmes:** Tesco · Kerry · Bank of Ireland · "
        "Big 4 · Diageo · Infineon"
    )

    st.divider()
    st.markdown(f"**Model:** `{GENERATION_MODEL}`")

    st.divider()
    if st.button("🗑️ Reset", use_container_width=True):
        st.session_state["tailored"] = None
        st.session_state["saved_jd"] = ""
        st.session_state["saved_base"] = PRAGHYA_BASE_RESUME
        st.session_state["cover_letter"] = None
        st.rerun()

    st.caption("Astra v2.1 — Praghya")

# Main UI
if not st.session_state["tailored"]:
    st.markdown(f"<h1 style='text-align: center;'>{PAGE_TITLE}</h1>", unsafe_allow_html=True)
    st.markdown(
        "<p style='text-align: center; color: #888;'>"
        "Paste a job description. Get a tailored, ATS-friendly resume.</p>",
        unsafe_allow_html=True,
    )
    st.divider()

    col_resume, col_jd = st.columns(2)

    with col_resume:
        st.subheader("📋 Base Resume")
        st.caption("Pre-filled with Praghya's profile. Edit only if you need a one-off tweak for this application.")
        base = st.text_area(
            "Base Resume",
            value=st.session_state["saved_base"],
            height=420,
            label_visibility="collapsed",
        )
        if st.button("↩️ Restore default base resume", use_container_width=True):
            st.session_state["saved_base"] = PRAGHYA_BASE_RESUME
            st.rerun()

    with col_jd:
        st.subheader("💼 Job Description")
        st.caption("Paste the full JD here.")
        jd = st.text_area(
            "Job Description",
            value=st.session_state["saved_jd"],
            height=420,
            label_visibility="collapsed",
            placeholder="Paste the full JD here…",
        )

    if st.button("✨ Generate Tailored Resume", type="primary", use_container_width=True):
        if not api_key:
            st.error("Need a Google API key to generate.")
        elif not jd.strip():
            st.warning("Please paste a job description.")
        elif not base.strip():
            st.warning("Base resume is empty. Click ‘Restore default base resume’ to bring it back.")
        else:
            st.session_state["saved_jd"] = jd
            st.session_state["saved_base"] = base
            with st.spinner("Tailoring resume to JD…"):
                model_out = call_gemini(api_key, jd, resume_text=base)
                if "error" in model_out:
                    st.error(model_out["error"])
                else:
                    final = assemble_resume(model_out)
                    st.session_state["tailored"] = final
                    st.rerun()

else:
    data = st.session_state["tailored"]

    # Top bar
    c1, c2 = st.columns([4, 1])
    with c1:
        st.markdown(f"## 🎯 Target: {data['target_company']}")
        st.caption(f"Tailored title: **{data['candidate_title']}**")
        archetype = data.get("jd_archetype", "")
        if archetype:
            archetype_display = {
                "supply_chain_coordinator": "🟢 Supply Chain Coordinator (strong fit)",
                "supply_chain_analyst_data_heavy": "🟡 Data-heavy SC Analyst (weak fit, see warning)",
                "graduate_programme_supply_chain": "🟢 Graduate Programme, Supply Chain (strong fit)",
                "graduate_programme_business_general": "🟢 Graduate Programme, Business/Management (strong fit)",
                "operations_executive_or_inventory": "🟢 Operations / Inventory (strong fit)",
                "pharma_supply_chain": "🟡 Pharma SC (transferable, frame retail as adaptable)",
                "unknown_stretch": "🔴 Stretch role, see warning",
            }.get(archetype, archetype)
            st.caption(f"JD archetype: {archetype_display}")
    with c2:
        if st.button("New JD", use_container_width=True):
            st.session_state["tailored"] = None
            st.session_state["saved_jd"] = ""
            st.session_state["cover_letter"] = None
            st.rerun()

    # JD fit warning — surface prominently if Astra flagged a mismatch
    fit_warning = data.get("jd_fit_warning", "").strip()
    if fit_warning:
        st.warning(f"⚠️ **JD fit warning:** {fit_warning}")

    # Tabs: Preview | Edit | Cover Letter | Download
    tab_preview, tab_edit, tab_cover, tab_download = st.tabs(
        ["👀 Preview", "📝 Edit", "✍️ Cover Letter", "📥 Download"]
    )

    with tab_preview:
        st.subheader("Professional Profile")
        st.write(data["summary"])

        st.subheader("Key Skills / Tools & Technologies")
        for cat, skill_list in data["skills"].items():
            st.markdown(f"- **{cat}:** {', '.join(skill_list)}")

        st.subheader("Professional Experience")
        for role in data["experience"]:
            st.markdown(
                f"**{role['role_title']}** | {role['company']} | "
                f"{role['location']} | {role['dates']}"
            )
            for r in role["responsibilities"]:
                st.markdown(f"- {r}")
            if role["achievements"]:
                st.markdown("**Achievements:**")
                for a in role["achievements"]:
                    st.markdown(f"- {a}")

        st.subheader("Education")
        for edu in data["education"]:
            st.markdown(f"**{edu['degree']}**")
            st.markdown(f"{edu['institution']} | {edu['dates']} | {edu['grade']}")
            if edu.get("extra"):
                st.write(edu["extra"])

        st.subheader("Certifications")
        for c in data["certifications"]:
            st.markdown(f"- {c}")

        st.subheader("Additional Information")
        for a in data["additional_info"]:
            st.markdown(f"- {a}")

    with tab_edit:
        with st.form("edit_form"):
            data["candidate_title"] = st.text_input("Title under name", data["candidate_title"])
            data["summary"] = st.text_area("Summary", data["summary"], height=160)

            st.markdown("##### Skills (comma-separated per category)")
            for cat in list(data["skills"].keys()):
                joined = ", ".join(data["skills"][cat])
                edited = st.text_area(cat, joined, height=70, key=f"sk_{cat}")
                data["skills"][cat] = [s.strip() for s in edited.split(",") if s.strip()]

            st.markdown("##### Experience")
            for i, role in enumerate(data["experience"]):
                with st.expander(f"{role['role_title']} @ {role['company']}", expanded=False):
                    resps_text = "\n".join(role["responsibilities"])
                    new_resps = st.text_area("Responsibilities (one per line)",
                                             resps_text, height=180, key=f"r_{i}")
                    role["responsibilities"] = [
                        line.strip() for line in new_resps.split("\n") if line.strip()
                    ]
                    if role["achievements"] or i == 0:  # show achievements box for Reliance
                        achs_text = "\n".join(role["achievements"])
                        new_achs = st.text_area("Achievements (one per line)",
                                                achs_text, height=80, key=f"a_{i}")
                        role["achievements"] = [
                            line.strip() for line in new_achs.split("\n") if line.strip()
                        ]

            if st.form_submit_button("💾 Save edits", type="primary"):
                st.session_state["tailored"] = data
                st.success("Saved.")
                st.rerun()

    with tab_cover:
        st.caption(
            "A short, human-sounding cover letter using the war story that best matches the JD. "
            "Generated separately so it doesn't slow down resume tailoring."
        )

        cover_btn_label = "✨ Draft Cover Letter" if not st.session_state["cover_letter"] else "🔄 Re-draft Cover Letter"
        if st.button(cover_btn_label, type="primary"):
            if not api_key:
                st.error("Need a Google API key to generate.")
            elif not st.session_state["saved_jd"].strip():
                st.warning("No saved JD found. Generate the resume first.")
            else:
                with st.spinner("Picking the right war story, drafting…"):
                    cl = generate_cover_letter(api_key, data, st.session_state["saved_jd"])
                    if cl.startswith("ERROR:"):
                        st.error(cl)
                    else:
                        st.session_state["cover_letter"] = cl
                        st.rerun()

        if st.session_state["cover_letter"]:
            edited = st.text_area(
                "Cover letter (editable)",
                st.session_state["cover_letter"],
                height=420,
            )
            # Persist any manual edits the user makes here
            st.session_state["cover_letter"] = edited

            try:
                cl_bytes = render_cover_letter_docx(
                    st.session_state["cover_letter"],
                    target_company=data.get("target_company", ""),
                )
                company_safe = re.sub(r"[^A-Za-z0-9_-]", "_",
                                      (data["target_company"] or "Company").strip()) or "Company"
                st.download_button(
                    "📄 Download Cover Letter (.docx)",
                    data=cl_bytes,
                    file_name=f"CoverLetter_Praghya_Prakhar_{company_safe}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                )
            except Exception as e:
                st.error(f"Cover letter render error: {e}")

    with tab_download:
        company = data["target_company"] or "Company"
        company_safe = re.sub(r"[^A-Za-z0-9_-]", "_", company.strip()) or "Company"
        filename_base = f"Praghya_Prakhar_{company_safe}"

        st.text_input("Filename (no extension)", filename_base, key="fname")
        fname = st.session_state.get("fname", filename_base)

        c1, c2 = st.columns(2)
        try:
            docx_bytes = render_docx(data)
            c1.download_button(
                "📄 Word (.docx)",
                data=docx_bytes,
                file_name=f"{fname}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )
        except Exception as e:
            c1.error(f"DOCX error: {e}")

        try:
            pdf_bytes = render_pdf(data)
            c2.download_button(
                "📕 PDF",
                data=pdf_bytes,
                file_name=f"{fname}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        except Exception as e:
            c2.error(f"PDF error: {e}")
